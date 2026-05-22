from __future__ import annotations

import argparse
import hashlib
import json
import os
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import time
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Literal, cast

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook, load_workbook
from pypdf import PdfReader


ROOT_DIR = Path(__file__).resolve().parents[4]
BASE_DIR = Path(__file__).resolve().parent
EVIDENCE_DIR = BASE_DIR / "evidence"
SUMMARY_PATH = EVIDENCE_DIR / "summary.json"
REPORT_PATH = BASE_DIR / "02-飞书附件100轮真实模型测试执行报告.md"
CASESET_PATH = BASE_DIR / "01-测试用例-飞书附件100轮真实模型场景.md"
TMP_PREFIX = "cycber_feishu_attach100_real_"
MODEL_PROXY_ENDPOINT = os.environ.get("CYCBER_REAL_MODEL_ENDPOINT", "http://127.0.0.1:8317/v1")
MODEL_VERIFY_TIMEOUT_SECONDS = float(os.environ.get("CYCBER_REAL_MODEL_VERIFY_TIMEOUT", "45"))


def _bootstrap_paths() -> None:
    paths = [
        ROOT_DIR / "apps" / "local-api",
        *ROOT_DIR.glob("packages/*"),
        *ROOT_DIR.glob("services/*"),
    ]
    for path in paths:
        if path.is_dir() and str(path) not in sys.path:
            sys.path.insert(0, str(path))


_bootstrap_paths()

from app.core.config import ChannelProviderSection  # noqa: E402
from app.main import create_app  # noqa: E402
from app.services.channel_connectors import FeishuMockConnector  # noqa: E402
from core_types import TaskArtifact  # noqa: E402


AttachmentKind = Literal["pdf", "docx", "txt", "xlsx"]
OutputKind = Literal["none", "txt", "docx", "xlsx", "pdf", "rename"]


@dataclass(frozen=True)
class AttachmentFixture:
    kind: AttachmentKind
    filename: str
    content_type: str
    path: Path
    key_terms: tuple[str, ...]
    sha256: str
    size_bytes: int


@dataclass(frozen=True)
class OperationSpec:
    slug: str
    title: str
    prompt: str
    expected_terms: tuple[str, ...]
    output_kind: OutputKind = "none"
    requires_extra_attachment: bool = False
    min_reply_score: int = 70
    min_file_score: int = 0


@dataclass(frozen=True)
class CaseSpec:
    case_id: str
    category: str
    title: str
    peer_ref: str
    prompt: str
    attachments: tuple[AttachmentFixture, ...]
    operation: OperationSpec


@dataclass
class CaseResult:
    case_id: str
    category: str
    title: str
    peer_ref: str
    prompt: str
    attachment_names: list[str]
    verdict: str
    notes: list[str]
    reply_text: str
    reply_score: int = 0
    file_score: int = 0
    turn_id: str | None = None
    conversation_id: str | None = None
    trace_id: str | None = None
    route_brain_id: str | None = None
    model_started: bool = False
    model_completed: bool = False
    usage_total_tokens: int | None = None
    attachment_count_seen: int = 0
    task_id: str | None = None
    task_status: str | None = None
    artifacts: list[dict[str, Any]] = field(default_factory=list)
    delivery_sent: bool = False
    files_sent: list[dict[str, Any]] = field(default_factory=list)


class ScenarioFeishuConnector(FeishuMockConnector):
    provider = "feishu"

    def __init__(self) -> None:
        super().__init__(ChannelProviderSection(enabled=True, poll_enabled=True))
        self.sent_text: list[dict[str, Any]] = []
        self.sent_file: list[dict[str, Any]] = []
        self._blobs: dict[str, bytes] = {}

    def register_blob(self, key: str, content: bytes) -> None:
        self._blobs[key] = content

    def send_count(self) -> int:
        return len(self.sent_text) + len(self.sent_file)

    async def send_text(
        self,
        *,
        provider_state_ref: str | None,
        provider_state: dict[str, Any] | None,
        recipient: str,
        text: str,
    ) -> Any:
        self.sent_text.append({"recipient": recipient, "text": text})
        return await super().send_text(
            provider_state_ref=provider_state_ref,
            provider_state=provider_state,
            recipient=recipient,
            text=text,
        )

    async def send_file(
        self,
        *,
        provider_state_ref: str | None,
        provider_state: dict[str, Any] | None,
        recipient: str,
        local_path: Path,
        content_type: str | None = None,
        filename: str | None = None,
    ) -> Any:
        self.sent_file.append(
            {
                "recipient": recipient,
                "path": str(local_path),
                "content_type": content_type,
                "filename": filename or local_path.name,
            }
        )
        return await super().send_file(
            provider_state_ref=provider_state_ref,
            provider_state=provider_state,
            recipient=recipient,
            local_path=local_path,
            content_type=content_type,
            filename=filename,
        )

    async def download_media(
        self,
        *,
        provider_state: dict[str, Any] | None,
        event: dict[str, Any],
        attachment: dict[str, Any],
    ) -> bytes:
        del provider_state, event
        key = str(attachment.get("file_key") or attachment.get("media_id") or "")
        if key in self._blobs:
            return self._blobs[key]
        return await super().download_media(
            provider_state=provider_state,
            event=event,
            attachment=attachment,
        )


def _hash_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _sha256_bytes(raw: bytes) -> str:
    return hashlib.sha256(raw).hexdigest()


def _copy_runtime_data() -> Path:
    temp_root = Path(tempfile.mkdtemp(prefix=TMP_PREFIX))
    data_dir = temp_root / "data"
    (data_dir / "sqlite").mkdir(parents=True, exist_ok=True)
    source_db = ROOT_DIR / "data" / "sqlite" / "app.db"
    if not source_db.exists():
        raise RuntimeError(f"source database not found: {source_db}")
    shutil.copy2(source_db, data_dir / "sqlite" / "app.db")
    source_secrets = ROOT_DIR / "data" / "secrets"
    if source_secrets.exists():
        shutil.copytree(source_secrets, data_dir / "secrets", dirs_exist_ok=True)
    for name in ["traces", "artifacts", "channel-providers", "backups", "restore-workspaces"]:
        (data_dir / name).mkdir(parents=True, exist_ok=True)
    _patch_brain_endpoint(data_dir / "sqlite" / "app.db")
    return data_dir


def _patch_brain_endpoint(db_path: Path) -> None:
    with sqlite3.connect(db_path) as conn:
        conn.execute(
            """
            UPDATE brains
            SET endpoint = ?, status = 'healthy', last_error_code = NULL, last_error_message = NULL
            WHERE provider IN ('openai_compatible', 'custom_openai_compatible')
               OR brain_id = 'brain_not_configured'
            """,
            (MODEL_PROXY_ENDPOINT,),
        )
        conn.commit()


def _fixture_dir(data_dir: Path) -> Path:
    path = data_dir / "test-fixtures" / "feishu-attachments-100"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _make_fixtures(data_dir: Path) -> dict[AttachmentKind, AttachmentFixture]:
    target = _fixture_dir(data_dir)
    terms = ("青藤计划", "12800", "Beta供应商", "6月15日", "陈澈")
    fixtures: dict[AttachmentKind, AttachmentFixture] = {}
    txt_path = target / "qingting-plan.txt"
    txt_path.write_text(
        "\n".join(
            [
                "青藤计划附件摘要",
                "预算合计 12800 元。",
                "关键风险：Beta供应商交付延期。",
                "截止日期：6月15日。",
                "负责人：陈澈。",
            ]
        ),
        encoding="utf-8",
    )
    fixtures["txt"] = _fixture("txt", txt_path, "text/plain", terms)

    docx_path = target / "qingting-plan.docx"
    document = Document()
    document.add_heading("青藤计划项目材料", level=1)
    for line in [
        "预算合计 12800 元。",
        "关键风险：Beta供应商交付延期。",
        "截止日期：6月15日。",
        "负责人：陈澈。",
    ]:
        document.add_paragraph(line)
    document.save(docx_path)
    fixtures["docx"] = _fixture(
        "docx",
        docx_path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        terms,
    )

    xlsx_path = target / "qingting-plan.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Data"
    sheet.append(["项目", "预算", "风险", "截止日期", "负责人"])
    sheet.append(["青藤计划", 12800, "Beta供应商交付延期", "6月15日", "陈澈"])
    workbook.save(xlsx_path)
    fixtures["xlsx"] = _fixture(
        "xlsx",
        xlsx_path,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        terms,
    )

    pdf_path = target / "qingting-plan.pdf"
    pdf_path.write_bytes(_minimal_pdf_bytes())
    fixtures["pdf"] = _fixture("pdf", pdf_path, "application/pdf", terms)
    return fixtures


def _fixture(
    kind: AttachmentKind,
    path: Path,
    content_type: str,
    key_terms: tuple[str, ...],
) -> AttachmentFixture:
    raw = path.read_bytes()
    return AttachmentFixture(
        kind=kind,
        filename=path.name,
        content_type=content_type,
        path=path,
        key_terms=key_terms,
        sha256=_sha256_bytes(raw),
        size_bytes=len(raw),
    )


def _minimal_pdf_bytes() -> bytes:
    stream = (
        "BT /F1 13 Tf 72 730 Td (Qingteng Plan) Tj "
        "0 -22 Td (Budget 12800 CNY) Tj "
        "0 -22 Td (Risk Beta supplier delay) Tj "
        "0 -22 Td (Deadline June 15) Tj "
        "0 -22 Td (Owner Chen Che) Tj ET"
    )
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n",
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
        f"5 0 obj << /Length {len(stream.encode('ascii'))} >> stream\n{stream}\nendstream endobj\n".encode(
            "ascii"
        ),
    ]
    chunks = [b"%PDF-1.4\n"]
    offsets = [0]
    for obj in objects:
        offsets.append(sum(len(chunk) for chunk in chunks))
        chunks.append(obj)
    xref_offset = sum(len(chunk) for chunk in chunks)
    xref = [b"xref\n0 6\n", b"0000000000 65535 f \n"]
    xref.extend(f"{offset:010d} 00000 n \n".encode("ascii") for offset in offsets[1:])
    chunks.extend(xref)
    chunks.append(b"trailer << /Size 6 /Root 1 0 R >>\n")
    chunks.append(f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii"))
    return b"".join(chunks)


def _operations() -> list[OperationSpec]:
    base_terms = ("青藤计划", "12800", "Beta供应商", "6月15日", "陈澈")
    return [
        OperationSpec("summary", "总结", "请阅读附件并总结成三点，必须只基于附件。", base_terms[:3]),
        OperationSpec("key_points", "归纳重点", "请归纳附件里的关键事实，按事实、风险、负责人回答。", base_terms),
        OperationSpec("action_items", "提取行动项", "请从附件中提取行动项、负责人和截止日期。", ("陈澈", "6月15日")),
        OperationSpec("risks", "风险归纳", "请只归纳附件里的风险和缓解建议。", ("Beta供应商", "风险")),
        OperationSpec("rename", "重命名建议", "请按附件内容给出一个标准文件名，不要声称已经改名。", ("青藤计划", "6月15日"), "rename"),
        OperationSpec("edit_polish", "编辑润色", "请把附件内容改写成老板能转发的简短同步。", ("青藤计划", "风险", "下一步")),
        OperationSpec("merge", "合并附件", "请合并两个附件，去重后输出一版统一摘要。", base_terms[:4], "none", True),
        OperationSpec("compare", "对比附件", "请对比两个附件，列出一致信息和冲突/缺口。", ("一致", "缺口", "青藤计划"), "none", True),
        OperationSpec("docx_out", "生成Word", "请基于附件生成一份 Word 复盘文件。", base_terms[:3], "docx", False, 70, 70),
        OperationSpec("xlsx_out", "生成Excel", "请基于附件生成一份 Excel 跟踪表。", base_terms[:3], "xlsx", False, 70, 70),
        OperationSpec("txt_out", "生成TXT", "请基于附件生成一份 TXT 摘要文件。", base_terms[:3], "txt", False, 70, 70),
        OperationSpec("pdf_out", "生成PDF", "请基于附件生成一份 PDF 简报文件。", base_terms[:3], "pdf", False, 70, 60),
        OperationSpec("numbers", "提取数字", "请提取附件里的金额、日期和责任人。", ("12800", "6月15日", "陈澈")),
        OperationSpec("timeline", "整理时间线", "请把附件整理成时间线和下一步计划。", ("6月15日", "下一步")),
        OperationSpec("decision_log", "决策记录", "请从附件生成决策记录，包含依据、风险、owner。", ("依据", "Beta供应商", "陈澈")),
        OperationSpec("qa", "问答化", "请把附件内容整理成 5 个问答。", ("青藤计划", "12800")),
        OperationSpec("glossary", "术语表", "请从附件抽取术语表和简短解释。", ("青藤计划", "Beta供应商")),
        OperationSpec("translate", "英文摘要", "请把附件核心内容整理成英文摘要，并保留金额和日期。", ("Qingteng", "12800", "June")),
        OperationSpec("compress", "100字压缩", "请把附件压缩到 100 字以内，但保留金额、风险、日期。", ("12800", "Beta", "6月15日")),
        OperationSpec("outline", "扩展大纲", "请基于附件扩展成汇报大纲。", ("青藤计划", "预算", "风险")),
        OperationSpec("toc", "目录结构", "请给附件拟一个目录结构。", ("预算", "风险", "负责人")),
        OperationSpec("redact", "脱敏版", "请输出附件脱敏版摘要，保留业务事实但不要泄露无关敏感信息。", ("青藤计划", "12800")),
        OperationSpec("filename", "文件名规范", "请根据附件内容提出 3 个文件命名规范方案。", ("青藤计划", "方案")),
        OperationSpec("followup", "追问清单", "请根据附件列出还需要向我追问的 5 个问题。", ("预算", "供应商", "截止")),
        OperationSpec("evidence", "证据报告", "请输出附件处理证据报告，说明读到了什么、没做什么、产物状态。", ("附件", "证据", "青藤计划")),
    ]


def _build_cases(fixtures: dict[AttachmentKind, AttachmentFixture]) -> list[CaseSpec]:
    cases: list[CaseSpec] = []
    kinds: tuple[AttachmentKind, ...] = ("pdf", "docx", "txt", "xlsx")
    for kind in kinds:
        for operation in _operations():
            case_no = len(cases) + 1
            attachments = [fixtures[kind]]
            if operation.requires_extra_attachment:
                attachments.append(fixtures["txt" if kind != "txt" else "docx"])
            cases.append(
                CaseSpec(
                    case_id=f"FATT100-{case_no:03d}",
                    category=f"{kind}_{operation.slug}",
                    title=f"{kind.upper()} {operation.title}",
                    peer_ref=f"oc_fatt100_{kind}_{operation.slug}",
                    prompt=operation.prompt,
                    attachments=tuple(attachments),
                    operation=operation,
                )
            )
    if len(cases) != 100:
        raise RuntimeError(f"expected 100 cases, got {len(cases)}")
    return cases


def _bind_feishu(client: TestClient) -> dict[str, Any]:
    started = client.post(
        "/api/channels/bind-sessions",
        json={
            "provider": "feishu",
            "requested_by_member_id": "mem_xiaoyao",
            "display_name_hint": "飞书附件100轮真实模型测试机器人",
        },
    )
    if started.status_code != 200:
        raise RuntimeError(started.text)
    payload = started.json()
    callback = client.get(
        "/api/channels/inbound/feishu/bind-callback",
        params={
            "state": payload["bind_session_id"],
            "code": "feishu-attach100-oauth-code",
            "tenant_key": "tenant_feishu_attach100_secret",
            "open_id": "ou_feishu_attach100_secret",
        },
    )
    if callback.status_code != 200:
        raise RuntimeError(callback.text)
    finalized = client.post(f"/api/channels/bind-sessions/{payload['bind_session_id']}/finalize")
    if finalized.status_code != 200:
        raise RuntimeError(finalized.text)
    return cast(dict[str, Any], finalized.json()["account"])


def _install_fake_feishu(client: TestClient) -> ScenarioFeishuConnector:
    registry = cast(Any, client.app).state.registry
    fake = ScenarioFeishuConnector()
    registry.channel_binding_service.connector_registry()._connectors["feishu"] = fake
    return fake


def _text_event(event_id: str, chat_id: str, sender_id: str, text: str) -> dict[str, Any]:
    return {
        "schema": "2.0",
        "header": {
            "event_id": event_id,
            "event_type": "im.message.receive_v1",
            "create_time": "2026-05-21T13:00:00+08:00",
        },
        "event": {
            "sender": {"sender_id": {"open_id": sender_id}, "sender_type": "user"},
            "message": {
                "message_id": f"om_{event_id}",
                "chat_id": chat_id,
                "chat_type": "p2p",
                "message_type": "text",
                "content": json.dumps({"text": text}, ensure_ascii=False),
            },
        },
    }


def _file_event(
    *,
    event_id: str,
    chat_id: str,
    sender_id: str,
    text: str,
    fixture: AttachmentFixture,
    file_key: str,
) -> dict[str, Any]:
    return {
        "schema": "2.0",
        "header": {
            "event_id": event_id,
            "event_type": "im.message.receive_v1",
            "create_time": "2026-05-21T13:00:00+08:00",
        },
        "event": {
            "sender": {"sender_id": {"open_id": sender_id}, "sender_type": "user"},
            "message": {
                "message_id": f"om_{event_id}",
                "chat_id": chat_id,
                "chat_type": "p2p",
                "message_type": "file",
                "content": json.dumps(
                    {
                        "text": text,
                        "file_key": file_key,
                        "file_name": fixture.filename,
                        "content_type": fixture.content_type,
                        "size": fixture.size_bytes,
                    },
                    ensure_ascii=False,
                ),
            },
        },
    }


def _latest_binding(client: TestClient) -> dict[str, Any] | None:
    payload = client.get("/api/channels/delivery-bindings", params={"provider": "feishu", "limit": 1})
    if payload.status_code != 200:
        raise RuntimeError(payload.text)
    items = payload.json()["items"]
    return cast(dict[str, Any], items[0]) if items else None


def _turn_payload(client: TestClient, turn_id: str) -> dict[str, Any]:
    response = client.get(f"/api/chat/turns/{turn_id}")
    if response.status_code != 200:
        raise RuntimeError(response.text)
    return cast(dict[str, Any], response.json())


def _turn_events(client: TestClient, turn_id: str) -> list[dict[str, Any]]:
    response = client.get(f"/api/chat/turns/{turn_id}/events")
    if response.status_code != 200:
        raise RuntimeError(response.text)
    return cast(list[dict[str, Any]], response.json()["items"])


def _visible_reply(events: list[dict[str, Any]]) -> str:
    text = "".join(
        str(item.get("payload", {}).get("payload", {}).get("text", ""))
        for item in events
        if item["event_type"] == "response.delta"
    )
    if text:
        return text
    for item in reversed(events):
        if item["event_type"] != "response.completed":
            continue
        payload = item.get("payload", {}).get("payload", {})
        response_plan = payload.get("response_plan", {}) or {}
        plain = str(response_plan.get("plain_text") or response_plan.get("summary") or "")
        if plain:
            return plain
    return ""


def _completed_structured_payload(events: list[dict[str, Any]]) -> dict[str, Any]:
    for item in reversed(events):
        if item["event_type"] != "response.completed":
            continue
        payload = item.get("payload", {}).get("payload", {})
        response_plan = payload.get("response_plan", {}) or {}
        return cast(dict[str, Any], response_plan.get("structured_payload") or {})
    return {}


def _model_summary(events: list[dict[str, Any]]) -> tuple[bool, bool, int | None, str | None]:
    started = any(item["event_type"] == "model.started" for item in events)
    completed = False
    total_tokens = None
    brain_id = None
    for item in events:
        payload = item.get("payload", {}).get("payload", {})
        if item["event_type"] == "model.started":
            brain_id = str(payload.get("brain_id") or "") or brain_id
        if item["event_type"] == "model.completed":
            completed = True
            usage = payload.get("usage") or {}
            if usage.get("total_tokens") is not None:
                total_tokens = int(usage["total_tokens"])
    return started, completed, total_tokens, brain_id


def _attachment_count_seen(events: list[dict[str, Any]]) -> int:
    for item in events:
        if item["event_type"] != "content.normalized":
            continue
        payload = item.get("payload", {}).get("payload", {})
        summary = payload.get("normalized_summary") or {}
        return int(summary.get("attachment_count") or 0)
    return 0


def _wait_for_new_turn(client: TestClient, previous_turn_id: str | None, timeout: float = 240.0) -> str:
    deadline = time.monotonic() + timeout
    last_status = None
    while time.monotonic() < deadline:
        binding = _latest_binding(client)
        if binding:
            turn_id = str(binding["turn_id"])
            if turn_id != previous_turn_id:
                turn = _turn_payload(client, turn_id)
                last_status = turn.get("status")
                if str(last_status) in {"completed", "failed", "cancelled"}:
                    return turn_id
        time.sleep(0.2)
    raise RuntimeError(f"new feishu turn was not observed, last_status={last_status}")


def _ensure_peer(client: TestClient, fake: ScenarioFeishuConnector, peer_ref: str, paired: set[str]) -> None:
    if peer_ref in paired:
        return
    fake.enqueue_event(_text_event(f"evt-pair-{peer_ref}", peer_ref, "ou_sender", "你好"))
    response = client.post("/api/channels/providers/feishu/poll-once")
    if response.status_code != 200:
        raise RuntimeError(response.text)
    pairings = client.get(
        "/api/channels/pairing-requests",
        params={"provider": "feishu", "status": "pending"},
    )
    if pairings.status_code != 200:
        raise RuntimeError(pairings.text)
    items = pairings.json()["items"]
    if not items:
        raise RuntimeError(f"no pairing created for {peer_ref}")
    approved = client.post(
        f"/api/channels/pairing-requests/{items[0]['pairing_request_id']}/approve",
        json={"member_id": "mem_xiaoyao", "reason": "feishu attachment real model test"},
    )
    if approved.status_code != 200:
        raise RuntimeError(approved.text)
    paired.add(peer_ref)


def _send_case(
    client: TestClient,
    fake: ScenarioFeishuConnector,
    spec: CaseSpec,
    paired: set[str],
) -> CaseResult:
    notes: list[str] = []
    _ensure_peer(client, fake, spec.peer_ref, paired)
    previous = _latest_binding(client)
    previous_turn_id = str(previous["turn_id"]) if previous else None
    previous_send_count = fake.send_count()
    file_sent_start = len(fake.sent_file)
    for index, fixture in enumerate(spec.attachments, start=1):
        file_key = f"file_{spec.case_id}_{index}_{fixture.sha256[:10]}"
        fake.register_blob(file_key, fixture.path.read_bytes())
        event_id = f"evt-{spec.case_id}-{index}-{_hash_text(spec.prompt)[:8]}"
        prompt = spec.prompt if index == 1 else "补充附件，用于合并或对比；请和上一份附件一起处理。"
        fake.enqueue_event(
            _file_event(
                event_id=event_id,
                chat_id=spec.peer_ref,
                sender_id="ou_sender",
                text=prompt,
                fixture=fixture,
                file_key=file_key,
            )
        )
    routed = client.post("/api/channels/providers/feishu/poll-once")
    if routed.status_code != 200:
        return _failed_result(spec, [f"poll_failed:{routed.status_code}"], routed.text)
    try:
        turn_id = _wait_for_new_turn(client, previous_turn_id)
    except Exception as exc:
        return _failed_result(spec, [f"turn_wait_failed:{exc}"], "")
    for _ in range(4):
        delivered = client.post("/api/channels/providers/feishu/deliver-due")
        if delivered.status_code != 200:
            notes.append(f"deliver_failed:{delivered.status_code}")
        time.sleep(0.1)
    turn = _turn_payload(client, turn_id)
    events = _turn_events(client, turn_id)
    structured = _completed_structured_payload(events)
    reply = _visible_reply(events)
    model_started, model_completed, usage_total, brain_id = _model_summary(events)
    task_id, task_status = _task_fields(structured)
    artifacts = _task_artifacts(client, task_id) if task_id else []
    reply_score, reply_notes = _score_reply(spec, reply, model_started, model_completed)
    file_score, file_notes = _score_files(client, spec, artifacts, fake.sent_file[file_sent_start:])
    notes.extend(reply_notes)
    notes.extend(file_notes)
    if str(turn.get("status")) != "completed":
        notes.append(f"turn_status:{turn.get('status')}")
    attachment_count = _attachment_count_seen(events)
    if attachment_count < len(spec.attachments):
        notes.append(f"attachment_count_seen:{attachment_count}")
    delivery_sent = fake.send_count() > previous_send_count
    if not delivery_sent:
        notes.append("delivery_not_sent")
    verdict = _verdict(spec, notes, reply_score, file_score)
    return CaseResult(
        case_id=spec.case_id,
        category=spec.category,
        title=spec.title,
        peer_ref=spec.peer_ref,
        prompt=spec.prompt,
        attachment_names=[item.filename for item in spec.attachments],
        verdict=verdict,
        notes=notes,
        reply_text=reply,
        reply_score=reply_score,
        file_score=file_score,
        turn_id=turn_id,
        conversation_id=turn.get("conversation_id"),
        trace_id=turn.get("trace_id"),
        route_brain_id=brain_id,
        model_started=model_started,
        model_completed=model_completed,
        usage_total_tokens=usage_total,
        attachment_count_seen=attachment_count,
        task_id=task_id,
        task_status=task_status,
        artifacts=artifacts,
        delivery_sent=delivery_sent,
        files_sent=list(fake.sent_file[file_sent_start:]),
    )


def _failed_result(spec: CaseSpec, notes: list[str], reply: str) -> CaseResult:
    return CaseResult(
        case_id=spec.case_id,
        category=spec.category,
        title=spec.title,
        peer_ref=spec.peer_ref,
        prompt=spec.prompt,
        attachment_names=[item.filename for item in spec.attachments],
        verdict="fail",
        notes=notes,
        reply_text=reply,
    )


def _task_fields(structured: dict[str, Any]) -> tuple[str | None, str | None]:
    task_status = cast(dict[str, Any], structured.get("task_status") or {})
    task_id = task_status.get("task_id")
    status = task_status.get("status")
    return (str(task_id) if task_id else None, str(status) if status else None)


def _task_artifacts(client: TestClient, task_id: str | None) -> list[dict[str, Any]]:
    if not task_id:
        return []
    response = client.get(f"/api/tasks/{task_id}/artifacts")
    if response.status_code != 200:
        return []
    return cast(list[dict[str, Any]], response.json().get("items") or [])


def _artifact_path(client: TestClient, artifact_id: str) -> Path:
    registry = cast(Any, client.app).state.registry
    artifact = cast(Any, client).portal.call(registry.artifact_store.get_artifact, artifact_id)
    return registry.artifact_store.path_for_artifact(TaskArtifact(**artifact.model_dump(mode="json")))


def _score_reply(
    spec: CaseSpec,
    reply: str,
    model_started: bool,
    model_completed: bool,
) -> tuple[int, list[str]]:
    notes: list[str] = []
    score = 100
    if not reply.strip():
        return 0, ["empty_reply"]
    if not (model_started and model_completed):
        score -= 45
        notes.append("real_model_not_completed")
    lowered = reply.lower()
    forbidden = ["trace_id", "tool_call_id", "approval_id", "sk-", "<tool_call", "<minimax"]
    for term in forbidden:
        if term in lowered:
            score -= 25
            notes.append(f"internal_or_secret_leak:{term}")
    all_expected = set(spec.operation.expected_terms)
    all_expected.update(term for item in spec.attachments for term in item.key_terms[:2])
    missing = [term for term in sorted(all_expected) if term and term not in reply]
    if missing:
        score -= min(50, len(missing) * 12)
        notes.append("missing_attachment_terms:" + ",".join(missing[:5]))
    if spec.operation.output_kind == "rename" and any(term in reply for term in ["已重命名", "已经改名"]):
        score -= 30
        notes.append("rename_false_done")
    if "附件" not in reply and "文件" not in reply:
        score -= 10
        notes.append("attachment_not_acknowledged")
    return max(0, score), notes


def _score_files(
    client: TestClient,
    spec: CaseSpec,
    artifacts: list[dict[str, Any]],
    sent_files: list[dict[str, Any]],
) -> tuple[int, list[str]]:
    output_kind = spec.operation.output_kind
    if output_kind in {"none", "rename"}:
        return 100, []
    notes: list[str] = []
    score = 100
    if not artifacts and not sent_files:
        return 0, ["expected_output_file_missing"]
    matches = _matching_artifacts(artifacts, output_kind)
    if not matches:
        score -= 55
        notes.append(f"artifact_type_missing:{output_kind}")
    else:
        artifact_score, artifact_notes = _inspect_artifact(client, output_kind, matches[-1], spec)
        score = min(score, artifact_score)
        notes.extend(artifact_notes)
    if not sent_files:
        score -= 20
        notes.append("file_not_delivered_to_feishu")
    return max(0, score), notes


def _matching_artifacts(artifacts: list[dict[str, Any]], output_kind: OutputKind) -> list[dict[str, Any]]:
    markers = {
        "docx": ("wordprocessingml.document", ".docx"),
        "xlsx": ("spreadsheetml.sheet", ".xlsx"),
        "pdf": ("application/pdf", ".pdf"),
        "txt": ("text/plain", ".txt", ".md"),
    }.get(output_kind, ())
    matches = []
    for item in artifacts:
        haystack = f"{item.get('content_type') or ''} {item.get('display_name') or ''}".lower()
        if any(marker in haystack for marker in markers):
            matches.append(item)
    return matches


def _inspect_artifact(
    client: TestClient,
    output_kind: OutputKind,
    artifact: dict[str, Any],
    spec: CaseSpec,
) -> tuple[int, list[str]]:
    notes: list[str] = []
    score = 100
    try:
        path = _artifact_path(client, str(artifact["artifact_id"]))
    except Exception as exc:
        return 20, [f"artifact_path_failed:{exc}"]
    if not path.exists():
        return 20, ["artifact_file_missing_on_disk"]
    text = ""
    try:
        if output_kind == "docx":
            doc = Document(str(path))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            if len(doc.paragraphs) < 2:
                score -= 20
                notes.append("docx_too_few_paragraphs")
        elif output_kind == "xlsx":
            workbook = load_workbook(path)
            rows = []
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    rows.append(" ".join(str(cell) for cell in row if cell is not None))
            text = "\n".join(rows)
            if not rows:
                score -= 30
                notes.append("xlsx_empty")
        elif output_kind == "pdf":
            reader = PdfReader(str(path))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if not text.strip():
                score -= 35
                notes.append("pdf_text_not_extractable")
        elif output_kind == "txt":
            text = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        return 40, [f"artifact_inspect_failed:{exc}"]
    required = set(spec.attachments[0].key_terms[:3])
    missing = [term for term in sorted(required) if term not in text]
    if missing:
        score -= min(50, len(missing) * 18)
        notes.append("artifact_content_missing:" + ",".join(missing))
    return max(0, score), notes


def _verdict(spec: CaseSpec, notes: list[str], reply_score: int, file_score: int) -> str:
    hard_prefixes = (
        "poll_failed",
        "turn_wait_failed",
        "empty_reply",
        "real_model_not_completed",
        "turn_status:",
        "internal_or_secret_leak",
    )
    if any(any(note.startswith(prefix) for prefix in hard_prefixes) for note in notes):
        return "fail"
    if reply_score < spec.operation.min_reply_score:
        return "fail"
    if file_score < spec.operation.min_file_score:
        return "fail"
    return "pass" if not notes else "warn"


def _write_caseset(cases: list[CaseSpec]) -> None:
    lines = [
        "# 飞书附件 100 轮真实模型测试用例",
        "",
        "- 入口：飞书渠道 mock connector，经 `poll-once -> channel ingress -> chat turn -> deliver-due`。",
        "- 模型要求：真实大脑模型，逐轮检查 `model.started` 与 `model.completed`。",
        "- 附件类型：PDF、Word、TXT、Excel。",
        "- 附件 oracle：隐藏在真实附件内容中，评分要求回复和产物命中关键事实。",
        "",
    ]
    for case in cases:
        lines.extend(
            [
                f"## {case.case_id} {case.title}",
                f"- 飞书 peer: `{case.peer_ref}`",
                f"- 附件: {', '.join(item.filename for item in case.attachments)}",
                f"- 输入: {case.prompt}",
                f"- 预期产物: `{case.operation.output_kind}`",
                "",
            ]
        )
    CASESET_PATH.write_text("\n".join(lines), encoding="utf-8")


def _write_outputs(results: list[CaseResult], *, model_verify: dict[str, Any], cases: list[CaseSpec]) -> None:
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    _write_caseset(cases)
    passed = sum(1 for item in results if item.verdict == "pass")
    warned = sum(1 for item in results if item.verdict == "warn")
    failed = sum(1 for item in results if item.verdict == "fail")
    summary = {
        "run_label": "FATT100-REAL-20260521",
        "entry": "feishu_mock_channel",
        "real_model_required": True,
        "model_endpoint": MODEL_PROXY_ENDPOINT,
        "model_verify": {
            key: value
            for key, value in model_verify.items()
            if key not in {"message", "verify_capabilities"}
        },
        "total": len(results),
        "passed": passed,
        "warned": warned,
        "failed": failed,
        "reply_score_avg": _avg([item.reply_score for item in results]),
        "file_score_avg": _avg([item.file_score for item in results]),
        "results": [asdict(item) for item in results],
    }
    SUMMARY_PATH.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        "# 飞书附件 100 轮真实模型测试执行报告",
        "",
        "- 测试入口：飞书 mock connector，经 `poll-once -> channel ingress -> chat turn -> deliver-due`。",
        "- 模型要求：真实模型调用，检查 `model.started` 与 `model.completed`。",
        f"- 模型端点：`{MODEL_PROXY_ENDPOINT}`。",
        f"- 结果：{passed} pass / {warned} warn / {failed} fail。",
        f"- 平均回复评分：{summary['reply_score_avg']}。",
        f"- 平均文件评分：{summary['file_score_avg']}。",
        "",
        "## 模型预检",
        "",
        f"- 状态：`{summary['model_verify'].get('status')}`",
        f"- 错误码：`{summary['model_verify'].get('error_code') or '-'}`",
        "",
        "## 明细",
        "",
        "| Case | 分类 | 结论 | 回复分 | 文件分 | 模型 | 附件数 | 产物 | 备注 |",
        "|---|---|---:|---:|---:|---|---:|---:|---|",
    ]
    for item in results:
        model = "ok" if item.model_started and item.model_completed else "no"
        lines.append(
            "| {case} | {category} | {verdict} | {reply_score} | {file_score} | {model} | {attach} | {artifact_count} | {notes} |".format(
                case=item.case_id,
                category=item.category,
                verdict=item.verdict,
                reply_score=item.reply_score,
                file_score=item.file_score,
                model=model,
                attach=item.attachment_count_seen,
                artifact_count=len(item.artifacts),
                notes=", ".join(item.notes) or "-",
            )
        )
    if results:
        lines.extend(["", "## 样例回复摘录", ""])
        for item in results[:20]:
            preview = item.reply_text.replace("\n", " ")[:220]
            lines.append(f"- `{item.case_id}`: {preview}")
    REPORT_PATH.write_text("\n".join(lines), encoding="utf-8")


def _avg(values: list[int]) -> float | None:
    if not values:
        return None
    return round(sum(values) / len(values), 2)


def run(*, limit: int | None = None) -> list[CaseResult]:
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    data_dir = _copy_runtime_data()
    temp_root = data_dir.parent
    fixtures = _make_fixtures(data_dir)
    cases = _build_cases(fixtures)
    if limit is not None:
        cases = cases[:limit]
    _write_caseset(cases)
    old_env = {
        key: os.environ.get(key)
        for key in [
            "CYCBER_ROOT",
            "CYCBER_DATA_DIR",
            "CYCBER_BROWSER_EXECUTOR",
            "FEISHU_APP_ID",
            "FEISHU_APP_SECRET",
            "USERPROFILE",
            "HOME",
        ]
    }
    try:
        os.environ["CYCBER_ROOT"] = str(ROOT_DIR)
        os.environ["CYCBER_DATA_DIR"] = str(data_dir)
        os.environ["CYCBER_BROWSER_EXECUTOR"] = "http_fallback"
        os.environ["FEISHU_APP_ID"] = "feishu-attach100-real-app"
        os.environ["FEISHU_APP_SECRET"] = "feishu-attach100-real-secret"
        os.environ["USERPROFILE"] = str(data_dir / "home")
        os.environ["HOME"] = str(data_dir / "home")
        (data_dir / "home" / "Desktop").mkdir(parents=True, exist_ok=True)

        verify_payload = _verify_real_model_subprocess(data_dir)
        if verify_payload.get("status_code") != 200 or verify_payload.get("status") != "healthy":
            _write_outputs([], model_verify=verify_payload, cases=cases)
            raise RuntimeError(f"real model verify failed: {verify_payload}")

        with TestClient(create_app()) as client:
            _bind_feishu(client)
            fake = _install_fake_feishu(client)
            paired: set[str] = set()
            results = [_send_case(client, fake, case, paired) for case in cases]
            _write_outputs(results, model_verify=verify_payload, cases=cases)
            return results
    finally:
        for key, value in old_env.items():
            if value is None:
                os.environ.pop(key, None)
            else:
                os.environ[key] = value
        shutil.rmtree(temp_root, ignore_errors=True)


def _verify_real_model_subprocess(data_dir: Path) -> dict[str, Any]:
    env = os.environ.copy()
    env["CYCBER_ROOT"] = str(ROOT_DIR)
    env["CYCBER_DATA_DIR"] = str(data_dir)
    env["CYCBER_BROWSER_EXECUTOR"] = "http_fallback"
    try:
        completed = subprocess.run(
            [
                sys.executable,
                str(Path(__file__).resolve()),
                "--preflight-only",
            ],
            cwd=str(ROOT_DIR),
            env=env,
            text=True,
            capture_output=True,
            timeout=MODEL_VERIFY_TIMEOUT_SECONDS,
        )
    except subprocess.TimeoutExpired:
        return {
            "brain_id": "brain_not_configured",
            "status": "unhealthy",
            "status_code": 598,
            "error_code": "MODEL_VERIFY_TIMEOUT",
            "timeout_seconds": MODEL_VERIFY_TIMEOUT_SECONDS,
        }
    stdout = completed.stdout.strip()
    if stdout:
        try:
            payload = json.loads(stdout.splitlines()[-1])
        except json.JSONDecodeError:
            payload = {
                "status": "unhealthy",
                "status_code": completed.returncode,
                "error_code": "MODEL_VERIFY_BAD_JSON",
                "stdout_tail": stdout[-500:],
            }
    else:
        payload = {
            "status": "unhealthy",
            "status_code": completed.returncode,
            "error_code": "MODEL_VERIFY_NO_STDOUT",
        }
    if completed.returncode != 0 and payload.get("status") == "healthy":
        payload["status"] = "unhealthy"
        payload["error_code"] = "MODEL_VERIFY_PROCESS_FAILED"
    if completed.stderr.strip():
        payload["stderr_tail"] = completed.stderr.strip()[-500:]
    return cast(dict[str, Any], payload)


def _preflight_child() -> dict[str, Any]:
    try:
        with TestClient(create_app()) as client:
            response = client.post("/api/brains/brain_not_configured/verify")
            payload = (
                response.json()
                if response.headers.get("content-type", "").startswith("application/json")
                else {"text": response.text}
            )
            payload["status_code"] = response.status_code
            return cast(dict[str, Any], payload)
    except Exception as exc:
        return {
            "brain_id": "brain_not_configured",
            "status": "unhealthy",
            "status_code": 599,
            "error_code": "MODEL_VERIFY_EXCEPTION",
            "error_summary": str(exc),
        }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--limit", type=int, default=None)
    parser.add_argument("--preflight-only", action="store_true")
    args = parser.parse_args()
    if args.preflight_only:
        print(json.dumps(_preflight_child(), ensure_ascii=False))
        return
    results = run(limit=args.limit)
    failed = [item for item in results if item.verdict == "fail"]
    print(
        json.dumps(
            {
                "total": len(results),
                "passed": sum(1 for item in results if item.verdict == "pass"),
                "warned": sum(1 for item in results if item.verdict == "warn"),
                "failed": len(failed),
                "summary": str(SUMMARY_PATH),
                "report": str(REPORT_PATH),
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    if failed:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
